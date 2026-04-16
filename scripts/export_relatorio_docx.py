# -*- coding: utf-8 -*-
"""
Exporta RELATORIO_FINAL_Atividade2_ML.md para rene_estevam_deckers_atividade_2.docx
com formatação de texto corrido (estilo habitual em trabalhos académicos / ABNT).

Inclui:
  - Margens (por defeito: esquerda 3 cm; demais 2 cm)
  - Recuo de primeira linha 1,25 cm no texto corrido
  - Títulos (níveis 1–3) alinhados à esquerda
  - Fonte Times New Roman ou Arial 12 pt (--font)
  - Citações longas (bloco Markdown > ) com recuo ~4 cm
  - Secção \"5. Referências\": lista bibliográfica final (não são citações no corpo do texto);
    aplica-se sangria pendente (uso comum na ABNT)
  - Legendas *Figura …* / *Tabela …* centradas

Não gera: capa, sumário automático nem numeração de páginas no rodapé — só o corpo do relatório.

Requer: pip install python-docx
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / "RELATORIO_FINAL_Atividade2_ML.md"
OUT_PATH = ROOT / "rene_estevam_deckers_atividade_2.docx"

FONT_TIMES = "Times New Roman"
FONT_ARIAL = "Arial"

MARGIN_LEFT_CM = 3.0
MARGIN_OTHER_CM = 2.0
FIRST_LINE_CM = 1.25
QUOTE_INDENT_CM = 4.0
REF_HANGING_CM = 1.25


def set_run_font(run, size_pt: int = 12, font_name: str = FONT_TIMES) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    try:
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
        rFonts.set(qn("w:eastAsia"), font_name)
        rFonts.set(qn("w:cs"), font_name)
    except Exception:
        pass


def apply_section_margins(section, left_cm: float, other_cm: float) -> None:
    section.left_margin = Cm(left_cm)
    section.right_margin = Cm(other_cm)
    section.top_margin = Cm(other_cm)
    section.bottom_margin = Cm(other_cm)


def format_paragraph_body(
    p,
    font_name: str,
    first_line: bool = True,
) -> None:
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(FIRST_LINE_CM) if first_line else Cm(0)
    for run in p.runs:
        if run.font.name and "Courier" in (run.font.name or ""):
            continue
        set_run_font(run, 12, font_name)


def format_paragraph_blockquote(p, font_name: str) -> None:
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(QUOTE_INDENT_CM)
    p.paragraph_format.right_indent = Cm(QUOTE_INDENT_CM)
    p.paragraph_format.first_line_indent = Cm(0)
    for run in p.runs:
        set_run_font(run, 11, font_name)


def format_paragraph_reference_entry(p, font_name: str) -> None:
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(REF_HANGING_CM)
    p.paragraph_format.first_line_indent = Cm(-REF_HANGING_CM)
    for run in p.runs:
        set_run_font(run, 12, font_name)


def format_paragraph_caption(p, font_name: str) -> None:
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)
    for run in p.runs:
        set_run_font(run, 12, font_name)


def format_list_item(p, font_name: str) -> None:
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Cm(0)
    for run in p.runs:
        set_run_font(run, 12, font_name)


def format_heading_left(p, font_name: str, size_pt: int) -> None:
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        set_run_font(run, size_pt, font_name)


def preprocess_links(text: str) -> str:
    return re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)


def _emit_run(
    paragraph,
    chunk: str,
    size_pt: int,
    font_name: str,
    *,
    bold: bool = False,
    italic: bool = False,
    code: bool = False,
) -> None:
    if not chunk:
        return
    r = paragraph.add_run(chunk)
    if code:
        r.font.name = "Courier New"
        r.font.size = Pt(size_pt - 1)
    r.bold = bold
    r.italic = italic
    if not code:
        set_run_font(r, size_pt, font_name)
    elif bold or italic:
        # código com contexto raro: aplica só peso se necessário
        r.bold = bold
        r.italic = italic


def add_inline_runs(
    paragraph,
    text: str,
    size_pt: int,
    font_name: str,
    *,
    bold: bool = False,
    italic: bool = False,
) -> None:
    """
    Markdown inline: `` `código` ``, **negrito**, *itálico* (inclui *aninhado* dentro de **negrito**).
    Ordem: links, depois `, depois **, depois *.
    """
    text = preprocess_links(text)
    i = 0
    n = len(text)
    while i < n:
        if text[i] == "`":
            j = text.find("`", i + 1)
            if j == -1:
                _emit_run(paragraph, text[i : i + 1], size_pt, font_name, bold=bold, italic=italic)
                i += 1
                continue
            chunk = text[i + 1 : j]
            _emit_run(
                paragraph,
                chunk,
                size_pt,
                font_name,
                bold=bold,
                italic=italic,
                code=True,
            )
            i = j + 1
            continue

        if text.startswith("**", i):
            j = text.find("**", i + 2)
            if j == -1:
                _emit_run(paragraph, text[i : i + 2], size_pt, font_name, bold=bold, italic=italic)
                i += 2
                continue
            inner = text[i + 2 : j]
            add_inline_runs(
                paragraph,
                inner,
                size_pt,
                font_name,
                bold=True,
                italic=italic,
            )
            i = j + 2
            continue

        if text[i] == "*":
            j = text.find("*", i + 1)
            if j == -1:
                _emit_run(paragraph, text[i : i + 1], size_pt, font_name, bold=bold, italic=italic)
                i += 1
                continue
            inner = text[i + 1 : j]
            add_inline_runs(
                paragraph,
                inner,
                size_pt,
                font_name,
                bold=bold,
                italic=True,
            )
            i = j + 1
            continue

        start = i
        while i < n and text[i] not in "`*":
            i += 1
        if i > start:
            _emit_run(
                paragraph,
                text[start:i],
                size_pt,
                font_name,
                bold=bold,
                italic=italic,
            )


def is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.count("|") >= 2


def is_table_sep(line: str) -> bool:
    s = line.strip().replace(" ", "")
    return bool(re.match(r"^\|?[-:|]+\|?$", s))


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and is_table_row(lines[i]):
        row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if is_table_sep(lines[i]):
            i += 1
            continue
        rows.append(row)
        i += 1
    return rows, i


def is_caption_line(stripped: str) -> bool:
    inner = stripped.strip()
    if len(inner) >= 2 and inner.startswith("*") and inner.endswith("*"):
        inner = inner[1:-1].strip()
    else:
        return False
    return bool(re.match(r"^(Figura|Tabela)\s+\d+", inner, re.I))


def export_main(
    md_path: Path,
    out_path: Path,
    font_name: str,
    margin_left: float,
    margin_other: float,
) -> None:
    if not md_path.is_file():
        print("Arquivo não encontrado:", md_path, file=sys.stderr)
        sys.exit(1)

    raw_lines = md_path.read_text(encoding="utf-8").splitlines()
    lines = [ln.rstrip() for ln in raw_lines]

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(12)

    section = doc.sections[0]
    apply_section_margins(section, margin_left, margin_other)

    md_dir = md_path.parent
    i = 0
    in_references = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("*[") and "Inserir figura" in stripped:
            i += 1
            continue

        if stripped.startswith("![") and "](" in stripped:
            m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
            if m:
                rel = m.group(2).strip()
                img_path = (md_dir / rel).resolve()
                if img_path.is_file():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.first_line_indent = Cm(0)
                    run = p.add_run()
                    run.add_picture(str(img_path), width=Cm(15))
                    p.paragraph_format.space_after = Pt(6)
                else:
                    p = doc.add_paragraph()
                    add_inline_runs(p, f"[Imagem não encontrada: {rel}]", 12, font_name)
                    format_paragraph_body(p, font_name, first_line=not in_references)
            i += 1
            continue

        if stripped.startswith("# "):
            t = stripped[2:].strip()
            p = doc.add_heading(t, level=1)
            format_heading_left(p, font_name, 14)
            i += 1
            continue

        if stripped.startswith("## "):
            t = stripped[3:].strip()
            in_references = bool(re.match(r"^5\.\s*Referências\b", t, re.I))
            p = doc.add_heading(t, level=2)
            format_heading_left(p, font_name, 13)
            i += 1
            continue

        if stripped.startswith("### "):
            t = stripped[4:].strip()
            p = doc.add_heading(t, level=3)
            format_heading_left(p, font_name, 12)
            i += 1
            continue

        if stripped.startswith("> "):
            inner = stripped[2:].strip()
            p = doc.add_paragraph()
            add_inline_runs(p, inner, 11, font_name)
            format_paragraph_blockquote(p, font_name)
            i += 1
            continue

        if stripped.startswith("- ") or (
            stripped.startswith("* ") and len(stripped) > 2
        ):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, stripped[2:].strip(), 12, font_name)
            format_list_item(p, font_name)
            i += 1
            continue

        if len(stripped) >= 2 and stripped.startswith("*") and stripped.endswith("*"):
            inner = stripped[1:-1]
            if inner and not inner.startswith("*"):
                if is_caption_line(stripped):
                    p = doc.add_paragraph()
                    r = p.add_run(inner)
                    r.italic = True
                    set_run_font(r, 12, font_name)
                    format_paragraph_caption(p, font_name)
                else:
                    p = doc.add_paragraph()
                    r = p.add_run(inner)
                    r.italic = True
                    set_run_font(r, 12, font_name)
                    if in_references:
                        format_paragraph_reference_entry(p, font_name)
                    else:
                        format_paragraph_body(p, font_name, first_line=True)
                i += 1
                continue

        if is_table_row(line):
            table_rows, new_i = parse_table(lines, i)
            if not table_rows:
                i = new_i
                continue
            cols = len(table_rows[0])
            table = doc.add_table(rows=len(table_rows), cols=cols)
            table.style = "Table Grid"
            for ri, row_cells in enumerate(table_rows):
                for ci, cell_text in enumerate(row_cells):
                    if ci < len(table.rows[ri].cells):
                        cell = table.rows[ri].cells[ci]
                        cell.text = cell_text
                        for para in cell.paragraphs:
                            if in_references:
                                format_paragraph_reference_entry(para, font_name)
                            else:
                                format_paragraph_body(para, font_name, first_line=False)
            doc.add_paragraph()
            i = new_i
            continue

        p = doc.add_paragraph()
        add_inline_runs(p, stripped, 12, font_name)
        if in_references:
            format_paragraph_reference_entry(p, font_name)
        else:
            format_paragraph_body(p, font_name, first_line=True)
        i += 1

    doc.save(out_path)
    print("Gerado:", out_path)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Exporta o relatório Markdown para .docx (formatação do texto; sem capa, sumário nem rodapé)."
    )
    parser.add_argument(
        "--font",
        choices=("times", "arial"),
        default="times",
        help="Fonte (por defeito: Times New Roman)",
    )
    parser.add_argument(
        "--margin-left",
        type=float,
        default=MARGIN_LEFT_CM,
        metavar="CM",
        help=f"Margem esquerda em cm (por defeito: {MARGIN_LEFT_CM})",
    )
    parser.add_argument(
        "--margin-other",
        type=float,
        default=MARGIN_OTHER_CM,
        metavar="CM",
        help=f"Margens superior, inferior e direita em cm (por defeito: {MARGIN_OTHER_CM})",
    )
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        default=OUT_PATH,
        help="Caminho do .docx de saída",
    )
    parser.add_argument(
        "-i",
        "--input",
        type=Path,
        default=MD_PATH,
        help="Caminho do .md de entrada",
    )
    args = parser.parse_args()
    font = FONT_ARIAL if args.font == "arial" else FONT_TIMES
    export_main(
        args.input,
        args.output,
        font_name=font,
        margin_left=args.margin_left,
        margin_other=args.margin_other,
    )


if __name__ == "__main__":
    main()
